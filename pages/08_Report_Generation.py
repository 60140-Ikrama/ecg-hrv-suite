"""Dashboard (CLO3) — Professional Report Generation with full metrics export"""
import streamlit as st
import pandas as pd
import io
from datetime import datetime
import sys, os
import numpy as np
import plotly.graph_objects as go

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from components.theme import (inject_stitch_theme, sentinel_header,
                               pipeline_status_bar, section_header, COLORS, 
                               PLOTLY_LAYOUT, set_layout, save_all_figures)
from components.sidebar_settings import render_sidebar_settings
from utils.hrv_analysis import interpret_hrv, detrended_fluctuation_analysis
from utils.heart_disease_detection import classify_cardiovascular_risk

st.set_page_config(page_title="Report Generation · Clinical Sentinel",
                   page_icon="📑", layout="wide")


# ── report builder ────────────────────────────────────────────────────────────

def build_latex_report(metrics_dict: dict, settings: dict, sqi_cache: dict) -> str:
    """Generates a professional LaTeX source for the analysis report."""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    latex = [
        "\\documentclass[a4paper,10pt]{article}",
        "\\usepackage[utf8]{inputenc}",
        "\\usepackage[margin=1in]{geometry}",
        "\\usepackage{booktabs}",
        "\\usepackage{graphicx}",
        "\\usepackage{color}",
        "\\usepackage{hyperref}",
        "",
        "\\title{Clinical Sentinel: ECG \\& HRV Analysis Report}",
        "\\author{Automated Research Suite v2.2.0}",
        f"\\date{{{ts}}}",
        "",
        "\\begin{document}",
        "\\maketitle",
        "",
        "\\section{Methodology}",
        "This report presents an automated cardiovascular assessment based on Signal Quality-Aware Adaptive Analysis.",
        "\\begin{itemize}",
        f"  \\item \\textbf{{Sampling Rate:}} {settings.get('sfreq', 250):.0f} Hz",
        f"  \\item \\textbf{{Filter Range:}} {settings.get('lowcut', 0.5):.2f} -- {settings.get('highcut', 40):.0f} Hz",
        f"  \\item \\textbf{{R-Peak Algorithm:}} {settings.get('rpeak_method', 'NeuroKit')}",
        "\\end{itemize}",
        ""
    ]

    for fname, m in metrics_dict.items():
        fname_esc = fname.replace("_", "\\_")
        sqi = sqi_cache.get(fname, {})
        conf = m.get("Confidence (%)", "N/A")
        
        latex += [
            f"\\section{{Analysis for File: {fname_esc}}}",
            f"\\textbf{{Signal Quality Index (SQI):}} {sqi.get('overall_sqi', 'N/A')}\\% ({sqi.get('quality_label', 'N/A')})\\\\",
            f"\\textbf{{Analysis Confidence:}} {conf}\\%\\\\",
            "",
            "\\subsection{Time-Domain Metrics}",
            "\\begin{tabular}{ll}",
            "\\toprule",
            "Metric & Value \\\\",
            "\\midrule",
            f"Mean RR & {m.get('Mean RR (ms)', 'N/A')} ms \\\\",
            f"SDNN & {m.get('SDNN (ms)', 'N/A')} ms \\\\",
            f"RMSSD & {m.get('RMSSD (ms)', 'N/A')} ms \\\\",
            f"Mean HR & {m.get('Mean HR (bpm)', 'N/A')} bpm \\\\",
            "\\bottomrule",
            "\\end{tabular}",
            ""
        ]

    latex += ["\\end{document}"]
    return "\n".join(latex)

def build_markdown_report(metrics_dict: dict, settings: dict,
                           sqi_cache: dict) -> str:
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    lines = []
    lines += [
        "# Clinical Sentinel — ECG & HRV Analysis Report",
        f"\n**Generated:** {ts}  |  "
        f"**Version:** Research-Grade Suite v2.2.0 (SQI-Aware)\n",
        "---\n",
    ]

    # ── 1. Methodology ────────────────────────────────────────────────────────
    lines += ["## 1. Methodology\n"]
    lines += [
        "### Signal Acquisition",
        f"- Files analysed: {', '.join(metrics_dict.keys())}",
        f"- Sampling frequency: **{settings.get('sfreq', 250):.0f} Hz**\n",
        "### Preprocessing",
        f"- Bandpass filter: **{settings.get('lowcut', 0.5):.2f}–"
        f"{settings.get('highcut', 40):.0f} Hz** "
        f"(Butterworth order {settings.get('filter_order', 4)})",
        f"- Baseline wander removal: "
        f"**{'Yes' if settings.get('remove_baseline', True) else 'No'}**",
        "### R-Peak Detection",
        f"- Algorithm: **{settings.get('rpeak_method', 'NeuroKit')}**  \n"
        "  *(Pan-Tompkins Custom implementation — CLO1)*\n",
        "### Ectopic Beat Correction",
        f"- Enabled: **{'Yes' if settings.get('remove_ectopic', True) else 'No'}**",
        f"- Detection method: **{settings.get('ectopic_method', 'median')}**",
        f"- Detection threshold: **{settings.get('ectopic_threshold', 20)}%**\n",
        "### HRV Analysis",
        f"- LF band: **{settings.get('lf_min', 0.04):.2f}–{settings.get('lf_max', 0.15):.2f} Hz**",
        f"- HF band: **{settings.get('hf_min', 0.15):.2f}–{settings.get('hf_max', 0.40):.2f} Hz**",
        f"- PSD: **Welch's method**\n",
        "---\n",
    ]

    # ── 2. Results ────────────────────────────────────────────────────────────
    lines += ["## 2. Results\n"]
    for fname, m in metrics_dict.items():
        lines += [f"### File: `{fname}`\n"]
        lines += [
            "#### Time-Domain Metrics",
            "| Metric | Value |",
            "|---|---|",
        ]
        for k in ["Mean RR (ms)", "SDNN (ms)", "RMSSD (ms)", "NN50", "pNN50 (%)", "Mean HR (bpm)"]:
            v = m.get(k, "N/A")
            s = f"{v:.2f}" if isinstance(v, float) else str(v)
            lines.append(f"| {k} | {s} |")
        lines.append("")

        if "LF Power (ms²)" in m or "LF Power (ms2)" in m:
            lines += [
                "#### Frequency-Domain Metrics",
                "| Metric | Value |",
                "|---|---|",
            ]
            for k in ["LF Power (ms²)", "HF Power (ms²)", "LF/HF Ratio"]:
                k_lookup = k.replace("²", "2") if k not in m else k
                v = m.get(k_lookup, "N/A")
                s = f"{v:.3f}" if isinstance(v, float) else str(v)
                lines.append(f"| {k} | {s} |")
            lines.append("")

    # ── 3. Clinical Interpretation ────────────────────────────────────────────
    lines += ["---\n", "## 3. Clinical Interpretation\n"]
    for fname, m in metrics_dict.items():
        interp = interpret_hrv(m, m)
        lines += [
            f"### `{fname}`",
            f"- **Overall HRV (SDNN):** {interp.get('sdnn_class','N/A')}",
            f"- **Vagal Tone (RMSSD):** {interp.get('autonomic','N/A')} (Reflects HF/Parasympathetic activity)",
            f"- **Sympathovagal Balance:** {interp.get('lf_hf_class','N/A')} (Reflects LF/HF Ratio)\n",
        ]

    lines += [
        "---\n",
        "*Report generated by Clinical Sentinel ECG & HRV Analysis Suite v2.0*",
    ]
    return "\n".join(lines)


# ── Chart & Document Generators ───────────────────────────────────────────────

def _generate_report_charts(filename: str) -> dict:
    """Generates the 8 mandatory charts for the report with robust error handling."""
    charts = {}
    errors = []
    sfreq = st.session_state.get("sfreq", 250.0)
    
    # Ensure trapezoid is available
    import numpy as np
    try:
        from numpy import trapezoid as _trapz
    except ImportError:
        _trapz = np.trapz
    
    def _export_fig(fig, key, width=800, height=350):
        try:
            # Try kaleido first (best for static PDFs)
            img = fig.to_image(format="png", width=width, height=height, engine="kaleido")
            charts[key] = img
        except Exception as e:
            try:
                # Fallback: let plotly decide engine
                img = fig.to_image(format="png", width=width, height=height)
                charts[key] = img
            except Exception as e2:
                errors.append(f"Chart '{key}' failed: {str(e2)}")

    # 1. ECG Signal (Raw vs Filtered)
    raw_sig = st.session_state.get("raw_signals", {}).get(filename)
    filt_sig = st.session_state.get("cleaned_signals", {}).get(filename)
    if filt_sig is not None:
        max_idx = min(len(filt_sig), int(5 * sfreq))
        t = np.arange(max_idx) / sfreq
        fig = go.Figure()
        if raw_sig is not None:
            fig.add_trace(go.Scatter(x=t, y=raw_sig[:max_idx], name="Raw", 
                                     line=dict(color="rgba(132,147,150,0.4)", width=1)))
        fig.add_trace(go.Scatter(x=t, y=filt_sig[:max_idx], name="Filtered", 
                                 line=dict(color="#00daf3", width=1.5)))
        set_layout(fig, "ECG Signal (Raw vs Filtered)")
        _export_fig(fig, "ecg_raw_filt")
        
    # 2. R-Peak Detection Overlay
    rpeaks = st.session_state.get("rpeaks", {}).get(filename)
    if filt_sig is not None and rpeaks is not None:
        max_idx = min(len(filt_sig), int(10 * sfreq))
        t = np.arange(max_idx) / sfreq
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=t, y=filt_sig[:max_idx], name="ECG", 
                                 line=dict(color="#00daf3", width=1)))
        rp_in = rpeaks[rpeaks < max_idx]
        fig.add_trace(go.Scatter(x=rp_in / sfreq, y=filt_sig[rp_in], mode="markers", 
                                 name="R-Peaks", marker=dict(color="#ff4b4b", size=10, symbol="triangle-up")))
        set_layout(fig, "R-Peak Detection Overlay")
        _export_fig(fig, "rpeaks")

    # 3. RR Tachogram
    clean_rr = st.session_state.get("clean_rr_intervals", {}).get(filename)
    if clean_rr is not None:
        fig = go.Figure()
        fig.add_trace(go.Scatter(y=clean_rr, mode='lines+markers', 
                                 line=dict(color="#00daf3"), name="RR"))
        set_layout(fig, "RR Tachogram", xaxis_title="Beat Index", yaxis_title="RR Interval (ms)")
        _export_fig(fig, "rr_tachogram")

    # 4. Ectopic Correction (Raw vs Clean RR)
    raw_rr = st.session_state.get("raw_rr_intervals", {}).get(filename)
    if raw_rr is not None and clean_rr is not None:
        fig = go.Figure()
        fig.add_trace(go.Scatter(y=raw_rr, name="Raw RR", line=dict(color="rgba(255,75,75,0.3)", width=1)))
        fig.add_trace(go.Scatter(y=clean_rr, name="Clean RR", line=dict(color="#c3f400")))
        set_layout(fig, "Ectopic Correction (Raw vs Clean RR)", xaxis_title="Beat Index", yaxis_title="RR (ms)")
        _export_fig(fig, "ectopic_corr")

    # 5. PSD (Frequency Domain)
    psd_raw = st.session_state.get("psd_data", {}).get(filename)
    if psd_raw is not None:
        freqs_p, psd_p = psd_raw
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=freqs_p, y=psd_p, fill="tozeroy", line=dict(color="#00daf3")))
        set_layout(fig, "Power Spectral Density (LF/HF)", xaxis_title="Frequency (Hz)", yaxis_title="Power (ms²/Hz)")
        _export_fig(fig, "psd")

    # 6. Poincare Plot
    if clean_rr is not None and len(clean_rr) > 2:
        rn, rn1 = clean_rr[:-1], clean_rr[1:]
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=rn, y=rn1, mode="markers", 
                                 marker=dict(color="#00daf3", size=4, opacity=0.5)))
        set_layout(fig, "Poincaré Plot (Non-linear)", xaxis_title="RR(n) ms", yaxis_title="RR(n+1) ms")
        _export_fig(fig, "poincare", width=500, height=500)

    # 7. DFA Plot
    if clean_rr is not None and len(clean_rr) > 32:
        dfa_res = detrended_fluctuation_analysis(clean_rr)
        scales, fluct = dfa_res.get("scales", []), dfa_res.get("fluct", [])
        if len(scales) > 3:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=np.log10(scales), y=np.log10(fluct), 
                                     mode='markers+lines', marker=dict(color="#00daf3")))
            a1 = dfa_res.get("alpha1")
            if a1 and np.isfinite(a1):
                mask = (scales >= 4) & (scales <= 16)
                if np.sum(mask) >= 2:
                    lx, ly = np.log10(scales[mask]), np.log10(fluct[mask])
                    c = np.polyfit(lx, ly, 1)
                    x_line = np.array([lx.min(), lx.max()])
                    fig.add_trace(go.Scatter(x=x_line, y=np.polyval(c, x_line), mode='lines', 
                                             name=f"α1={a1:.2f}", line=dict(color="#c3f400", dash='dash')))
            set_layout(fig, "Detrended Fluctuation Analysis (DFA)", xaxis_title="log10(Scale n)", yaxis_title="log10(F(n))")
            _export_fig(fig, "dfa")

    # 8. RR Histogram
    if clean_rr is not None and len(clean_rr) > 5:
        fig = go.Figure()
        fig.add_trace(go.Histogram(x=clean_rr, nbinsx=40, marker_color="#00daf3"))
        mean_rr = float(np.mean(clean_rr))
        std_rr  = float(np.std(clean_rr))
        for offset, lbl in [(-std_rr, "-1σ"), (std_rr, "+1σ")]:
            fig.add_vline(x=mean_rr + offset, line_dash="dot", line_color="#c3f400")
        set_layout(fig, "RR Interval Histogram", xaxis_title="RR Interval (ms)", yaxis_title="Count")
        _export_fig(fig, "rr_histogram", height=320)

    if errors:
        st.session_state["report_errors"] = errors
    return charts

def _safe(v):
    if isinstance(v, (float, np.float64, np.float32)): return f"{v:.3f}"
    return "N/A" if v is None else str(v)


def build_pdf_report(metrics_dict: dict, settings: dict, sqi_cache: dict) -> bytes:
    import io as _io
    import tempfile, os as _os
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors as rl_colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, Image as RLImage, PageBreak)
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Register Unicode font
    import pathlib as _pl
    FONT_NAME = "DejaVu"
    try:
        import matplotlib
        fdir = _pl.Path(matplotlib.get_data_path()) / "fonts" / "ttf"
        fpath_reg = fdir / "DejaVuSans.ttf"
        fpath_bold = fdir / "DejaVuSans-Bold.ttf"
        
        if fpath_reg.exists():
            pdfmetrics.registerFont(TTFont("DejaVu", str(fpath_reg)))
            if fpath_bold.exists():
                pdfmetrics.registerFont(TTFont("DejaVu-Bold", str(fpath_bold)))
                # Link regular and bold into a family
                from reportlab.pdfbase.pdfmetrics import registerFontFamily
                registerFontFamily("DejaVu", normal="DejaVu", bold="DejaVu-Bold")
            FONT_NAME = "DejaVu"
        else:
            FONT_NAME = "Helvetica" # Fallback
    except:
        FONT_NAME = "Helvetica"

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, margin=1.5*cm)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", fontName=FONT_NAME, fontSize=22, spaceAfter=14, textColor=rl_colors.HexColor("#003366"), alignment=1)
    h2 = ParagraphStyle("h2", fontName=FONT_NAME, fontSize=16, spaceAfter=10, textColor=rl_colors.HexColor("#1a5276"), spaceBefore=12)
    body = ParagraphStyle("body", fontName=FONT_NAME, fontSize=10, leading=12)

    story = [Paragraph("ECG & HRV Analysis Report", h1), Spacer(1, 1*cm)]
    
    for fname, m in metrics_dict.items():
        sqi = sqi_cache.get(fname, {})
        story.append(Paragraph(f"Analysis for File: {fname}", h2))
        
        # Section: Global Summary
        story.append(Paragraph("<b>Global Analysis Summary</b>", body))
        summary_data = [
            ["Metric", "Value", "Confidence / Details"],
            ["Mean HR", f"{m.get('Mean HR (bpm)', 'N/A')}", "Average Heart Rate"],
            ["SDNN", _safe(m.get('SDNN (ms)')), "Overall Variability"],
            ["RMSSD", _safe(m.get('RMSSD (ms)')), "Parasympathetic Activity"],
            ["Analysis Confidence", f"{m.get('Confidence (%)','N/A')}%", "Metric Reliability"],
            ["Signal Quality", sqi.get('quality_label', 'N/A'), f"SQI Score: {sqi.get('overall_sqi','N/A')}%"]
        ]
        t_sum = Table(summary_data, colWidths=[4.5*cm, 3.5*cm, 8*cm])
        t_sum.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, rl_colors.grey),
            ('BACKGROUND', (0,0), (-1,0), rl_colors.HexColor("#d5e8f7")),
            ('FONTNAME', (0,0), (-1,0), FONT_NAME + "-Bold" if FONT_NAME != "Helvetica" else "Helvetica-Bold"),
        ]))
        story.append(t_sum)
        story.append(Spacer(1, 0.6*cm))

        # Section: Time-Domain
        story.append(Paragraph("Time-Domain Metrics", body))
        data = [["Metric", "Value", "Unit"]]
        for k, unit in [("Mean RR", "ms"), ("SDNN", "ms"), ("RMSSD", "ms"), ("NN50", "beats"), ("pNN50", "%")]:
            val = m.get(k) or m.get(k+" (ms)") or m.get(k+" (%)") or m.get(k+" (beats)")
            data.append([k, _safe(val), unit])
        t = Table(data, colWidths=[6*cm, 5*cm, 5*cm])
        t.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, rl_colors.grey), ('BACKGROUND', (0,0), (-1,0), rl_colors.lightgrey)]))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))

        # Section: Frequency-Domain
        story.append(Paragraph("Frequency-Domain Metrics (PSD)", body))
        data_f = [["Metric", "Value", "Unit"]]
        for k, unit in [("VLF Power", "ms²"), ("LF Power", "ms²"), ("HF Power", "ms²"), ("LF/HF Ratio", "ratio"), ("Total Power", "ms²")]:
            val = m.get(k) or m.get(k+" (ms2)") or m.get(k+" (ms²)")
            data_f.append([k, _safe(val), unit])
        t_f = Table(data_f, colWidths=[6*cm, 5*cm, 5*cm])
        t_f.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, rl_colors.grey), ('BACKGROUND', (0,0), (-1,0), rl_colors.lightgrey)]))
        story.append(t_f)
        story.append(Spacer(1, 0.5*cm))

        # Section: Non-Linear
        story.append(Paragraph("Non-Linear HRV Metrics", body))
        data_nl = [["Metric", "Value", "Description"]]
        for k, desc in [("SD1", "Short-term variability"), ("SD2", "Long-term variability"), 
                        ("SD1/SD2", "Short/Long term ratio"), ("DFA α1", "Short-term fractal exponent"),
                        ("Sample Entropy", "Signal complexity")]:
            val = m.get(k) or m.get(k+" (ms)") or m.get(k+" (ms2)")
            data_nl.append([k, _safe(val), desc])
        t_nl = Table(data_nl, colWidths=[4*cm, 4*cm, 8*cm])
        t_nl.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, rl_colors.grey), ('BACKGROUND', (0,0), (-1,0), rl_colors.lightgrey)]))
        story.append(t_nl)
        story.append(Spacer(1, 0.5*cm))

        # Interpretation
        story.append(Paragraph("<b>Detailed Clinical Assessment</b>", body))
        interp = interpret_hrv(m, m)
        assess_text = (
            f"The analyzed signal for <b>{fname}</b> shows an {interp.get('sdnn_class','—')}. "
            f"Regarding vagal regulation, the findings indicate a {interp.get('autonomic','—').lower()}. "
            f"The sympathovagal balance, as measured by the LF/HF ratio, suggests {interp.get('lf_hf_class','—').lower()}."
        )
        story.append(Paragraph(assess_text, body))
        story.append(Spacer(1, 0.5*cm))

        # Physiology Reference Table
        story.append(Paragraph("<b>Physiology Reference Guide (Normal Resting Ranges)</b>", body))
        ref_data = [
            ["Metric", "Normal Range", "Clinical Relevance"],
            ["HR (Mean)", "60 - 100 bpm", "Resting heart rate"],
            ["SDNN", "50 - 100 ms", "Overall autonomic health"],
            ["RMSSD", "20 - 50 ms", "Vagal tone / recovery"],
            ["LF/HF Ratio", "1.0 - 2.0", "Sympathovagal balance"],
            ["DFA α1", "0.9 - 1.2", "Healthy fractal scaling"]
        ]
        t_ref = Table(ref_data, colWidths=[4*cm, 4*cm, 8*cm])
        t_ref.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, rl_colors.grey),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('BACKGROUND', (0,0), (-1,0), rl_colors.whitesmoke)
        ]))
        story.append(t_ref)
        story.append(Spacer(1, 1*cm))

        # MANDATORY GRAPHS
        charts = _generate_report_charts(fname)
        chart_order = [
            ("ecg_raw_filt",  "Figure 1: ECG Signal (Raw vs Filtered)"),
            ("rpeaks",        "Figure 2: R-Peak Detection Overlay"),
            ("rr_tachogram",  "Figure 3: RR Tachogram (Interval Variation)"),
            ("rr_histogram",  "Figure 4: RR Interval Histogram"),
            ("ectopic_corr",  "Figure 5: Ectopic Beat Correction"),
            ("psd",           "Figure 6: Power Spectral Density (LF/HF Bands)"),
            ("poincare",      "Figure 7: Poincaré Plot (Non-linear Scatter)"),
            ("dfa",           "Figure 8: Detrended Fluctuation Analysis (DFA)"),
        ]
        
        for ckey, clabel in chart_order:
            if ckey in charts and charts[ckey]:
                story.append(Paragraph(clabel, body))
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                    tmp.write(charts[ckey])
                    tpath = tmp.name
                # Adjust size for poincare and dfa to fit well
                c_height = 6*cm
                if ckey == "poincare": c_height = 10*cm
                story.append(RLImage(tpath, width=16*cm, height=c_height))
                _os.remove(tpath)
                story.append(Spacer(1, 0.4*cm))
            else:
                st.session_state.setdefault("report_errors", []).append(f"Chart '{ckey}' missing; skipped in PDF.")
        # ── Heart Disease Risk Section ──────────────────────────────────────
        story.append(Paragraph("<b>Heart Disease Risk Assessment</b>", body))
        raw_rr_cache = {}
        try:
            import streamlit as _st
            raw_rr_cache = _st.session_state.get("raw_rr_intervals", {})
        except Exception:
            pass
        raw_rr_f = raw_rr_cache.get(fname)
        pct_e = 0.0
        if raw_rr_f is not None and len(raw_rr_f) > 0:
            from utils.hrv_analysis import detect_ectopic_beats as _deb
            _msk = _deb(raw_rr_f)
            pct_e = float(np.sum(_msk)) / len(raw_rr_f) * 100
        risk_res = classify_cardiovascular_risk(m, pct_ectopic=pct_e, use_ml=False)
        risk_level = risk_res["risk_level"]
        risk_score = risk_res["score"]
        risk_conf  = risk_res["confidence"]
        risk_color_map = {"Normal": rl_colors.HexColor("#4caf7d"),
                          "Mild Risk": rl_colors.HexColor("#ffba38"),
                          "High Risk": rl_colors.HexColor("#f44336")}
        risk_para_style = ParagraphStyle(
            "risk", fontName=FONT_NAME, fontSize=13, spaceAfter=6,
            textColor=risk_color_map.get(risk_level, rl_colors.black), alignment=1)
        story.append(Paragraph(f"Risk Level: {risk_level}", risk_para_style))
        story.append(Paragraph(
            f"Risk Score: {risk_score:.0f}/100  |  Confidence: {risk_conf:.0f}%", body))
        story.append(Spacer(1, 0.3*cm))
        flag_rows = [["Metric", "Value", "Status", "Clinical Note"]]
        for metric, info in risk_res.get("flags", {}).items():
            flag_rows.append([
                metric, info.get("value", "N/A"),
                info.get("status", "N/A").replace("_", " ").title(),
                info.get("clinical_note", "")[:80],
            ])
        if len(flag_rows) > 1:
            t_risk = Table(flag_rows, colWidths=[4*cm, 2.5*cm, 2.5*cm, 7*cm])
            t_risk.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.4, rl_colors.grey),
                ('BACKGROUND', (0,0), (-1,0), rl_colors.HexColor("#d5e8f7")),
                ('FONTSIZE', (0,0), (-1,-1), 7),
            ]))
            story.append(t_risk)
        story.append(Spacer(1, 0.5*cm))

        story.append(PageBreak())

    # ── Multi-file comparison summary ──────────────────────────────────────────
    if len(metrics_dict) > 1:
        story.append(Paragraph("Multi-File Comparison Summary", h2))
        comp_keys = ["Mean HR (bpm)", "SDNN (ms)", "RMSSD (ms)",
                     "pNN50 (%)", "LF/HF Ratio", "SD1 (ms)", "SD2 (ms)"]
        comp_header = ["Metric"] + [f[:18] for f in metrics_dict]
        comp_rows   = [comp_header]
        for k in comp_keys:
            row = [k]
            for f, mv in metrics_dict.items():
                v = mv.get(k)
                row.append(_safe(v) if v is not None else "N/A")
            comp_rows.append(row)
        col_w = [4.5*cm] + [max(2*cm, 12*cm // len(metrics_dict))] * len(metrics_dict)
        t_comp = Table(comp_rows, colWidths=col_w)
        t_comp.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, rl_colors.grey),
            ('BACKGROUND', (0,0), (-1,0), rl_colors.HexColor("#d5e8f7")),
            ('FONTSIZE', (0,0), (-1,-1), 8),
        ]))
        story.append(t_comp)
        story.append(Spacer(1, 0.6*cm))

    doc.build(story)
    return buf.getvalue()


def build_html_report(metrics_dict: dict, settings: dict, sqi_cache: dict) -> str:
    """Generates a self-contained, premium HTML research report (MATLAB Publish style)."""
    import base64
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    html = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        "  <meta charset='UTF-8'>",
        "  <meta name='viewport' content='width=device-width, initial-scale=1.0'>",
        "  <title>Clinical Sentinel · ECG-HRV Research Report</title>",
        "  <link href='https://fonts.googleapis.com/css2?family=Inter:wght@400;600;800&family=Manrope:wght@700;800&family=Fira+Code&display=swap' rel='stylesheet'>",
        "  <style>",
        "    :root { --primary: #003366; --secondary: #1a5276; --accent: #c3f400; --bg: #f8fafc; --card: #ffffff; --text: #1e293b; --muted: #64748b; }",
        "    body { font-family: 'Inter', sans-serif; background: var(--bg); color: var(--text); line-height: 1.6; margin: 0; padding: 0; }",
        "    .container { max-width: 900px; margin: 40px auto; background: var(--card); box-shadow: 0 10px 25px rgba(0,0,0,0.05); border-radius: 12px; padding: 60px; }",
        "    header { border-bottom: 2px solid #edf2f7; padding-bottom: 30px; margin-bottom: 40px; }",
        "    h1 { font-family: 'Manrope'; font-weight: 800; font-size: 2.5rem; color: var(--primary); margin: 0; }",
        "    .meta { display: flex; justify-content: space-between; color: var(--muted); font-size: 0.9rem; margin-top: 15px; }",
        "    nav { position: sticky; top: 20px; float: left; width: 220px; margin-left: -280px; font-size: 0.85rem; }",
        "    nav ul { list-style: none; padding: 0; }",
        "    nav li { margin-bottom: 8px; }",
        "    nav a { color: var(--muted); text-decoration: none; transition: 0.2s; }",
        "    nav a:hover { color: var(--secondary); }",
        "    section { margin-bottom: 50px; scroll-margin-top: 40px; }",
        "    h2 { font-family: 'Manrope'; font-size: 1.5rem; color: var(--secondary); border-left: 4px solid var(--accent); padding-left: 15px; margin-bottom: 25px; }",
        "    h3 { font-size: 1.1rem; margin-top: 30px; color: var(--primary); }",
        "    table { width: 100%; border-collapse: collapse; margin: 20px 0; font-size: 0.9rem; }",
        "    th { background: #f1f5f9; text-align: left; padding: 12px; font-weight: 600; border-bottom: 2px solid #e2e8f0; }",
        "    td { padding: 12px; border-bottom: 1px solid #edf2f7; }",
        "    .metric-val { font-family: 'Fira Code', monospace; font-weight: 600; color: #0f172a; }",
        "    .chart-container { margin: 30px 0; text-align: center; background: #fafafa; border-radius: 8px; padding: 20px; border: 1px solid #f1f5f9; }",
        "    .chart-container img { max-width: 100%; border-radius: 4px; box-shadow: 0 4px 12px rgba(0,0,0,0.05); }",
        "    .caption { font-size: 0.8rem; color: var(--muted); margin-top: 10px; font-style: italic; }",
        "    .risk-banner { padding: 20px; border-radius: 8px; display: flex; align-items: center; gap: 20px; margin: 30px 0; }",
        "    footer { text-align: center; margin-top: 60px; padding-top: 30px; border-top: 1px solid #edf2f7; color: var(--muted); font-size: 0.8rem; }",
        "    @media (max-width: 1200px) { nav { display: none; } .container { margin: 20px; padding: 30px; } }",
        "  </style>",
        "</head>",
        "<body>",
        "  <div class='container'>",
        "    <nav><ul>",
        "      <li><a href='#methodology'>1. Methodology</a></li>",
        "      <li><a href='#results'>2. Analysis Results</a></li>",
        "      <li><a href='#visualization'>3. Visualization</a></li>",
        "      <li><a href='#interpretation'>4. Clinical Assessment</a></li>",
        "    </ul></nav>",
        "    <header>",
        "      <h1>ECG & HRV Research Report</h1>",
        "      <div class='meta'>",
        f"        <span>Automated Research Suite v2.2.0</span>",
        f"        <span>Date: {ts}</span>",
        "      </div>",
        "    </header>",
        "",
        "    <section id='methodology'>",
        "      <h2>1. Methodology</h2>",
        "      <p>This report presents an automated cardiovascular assessment using Signal Quality-Aware Adaptive Analysis. All signals were processed through a Butterworth bandpass filter and ectopic beat correction was applied where necessary.</p>",
        "      <table>",
        "        <tr><td>Sampling Frequency</td><td class='metric-val'>" + str(settings.get('sfreq', 250)) + " Hz</td></tr>",
        "        <tr><td>Bandpass Filter</td><td class='metric-val'>" + str(settings.get('lowcut', 0.5)) + " - " + str(settings.get('highcut', 40)) + " Hz</td></tr>",
        "        <tr><td>R-Peak Algorithm</td><td class='metric-val'>" + str(settings.get('rpeak_method', 'NeuroKit')) + "</td></tr>",
        "      </table>",
        "    </section>"
    ]

    for fname, m in metrics_dict.items():
        sqi = sqi_cache.get(fname, {})
        html += [
            f"<section id='results-{fname}'>",
            f"  <h2>2. Analysis for File: {fname}</h2>",
            f"  <p>Overall Signal Quality Index (SQI): <strong>{sqi.get('overall_sqi','N/A')}%</strong> ({sqi.get('quality_label','N/A')})</p>",
            "  <h3>Time-Domain Metrics</h3>",
            "  <table>",
            "    <thead><tr><th>Metric</th><th>Value</th><th>Unit</th><th>Description</th></tr></thead>",
            "    <tbody>",
            f"      <tr><td>Mean RR</td><td class='metric-val'>{_safe(m.get('Mean RR (ms)'))}</td><td>ms</td><td>Average beat-to-beat interval</td></tr>",
            f"      <tr><td>SDNN</td><td class='metric-val'>{_safe(m.get('SDNN (ms)'))}</td><td>ms</td><td>Overall heart rate variability</td></tr>",
            f"      <tr><td>RMSSD</td><td class='metric-val'>{_safe(m.get('RMSSD (ms)'))}</td><td>ms</td><td>Parasympathetic activity</td></tr>",
            f"      <tr><td>Mean HR</td><td class='metric-val'>{_safe(m.get('Mean HR (bpm)'))}</td><td>bpm</td><td>Average heart rate</td></tr>",
            "    </tbody>",
            "  </table>"
        ]

        # Embed Images
        html.append("<section id='visualization'>")
        html.append("  <h3>Diagnostic Visualizations</h3>")
        charts = _generate_report_charts(fname)
        chart_order = [
            ("ecg_raw_filt",  "Figure 1: ECG Signal (Raw vs Filtered)"),
            ("rpeaks",        "Figure 2: R-Peak Detection Overlay"),
            ("rr_tachogram",  "Figure 3: RR Tachogram (Interval Variation)"),
            ("psd",           "Figure 4: Power Spectral Density (LF/HF Bands)"),
            ("poincare",      "Figure 5: Poincaré Plot (Non-linear Scatter)"),
        ]
        
        for ckey, clabel in chart_order:
            if ckey in charts:
                b64 = base64.b64encode(charts[ckey]).decode()
                html += [
                    "  <div class='chart-container'>",
                    f"    <img src='data:image/png;base64,{b64}' alt='{clabel}'>",
                    f"    <div class='caption'>{clabel}</div>",
                    "  </div>"
                ]
        html.append("</section>")

        # Clinical Assessment & Risk
        interp = interpret_hrv(m, m)
        raw_rr_f = st.session_state.get("raw_rr_intervals", {}).get(fname)
        pct_e = 0.0
        if raw_rr_f is not None and len(raw_rr_f) > 0:
            from utils.hrv_analysis import detect_ectopic_beats as _deb
            _msk = _deb(raw_rr_f)
            pct_e = float(np.sum(_msk)) / len(raw_rr_f) * 100
        
        risk_res = classify_cardiovascular_risk(m, pct_ectopic=pct_e, use_ml=False)
        risk_level = risk_res["risk_level"]
        risk_color = {"Normal": "#c3f400", "Mild Risk": "#ffba38", "High Risk": "#ffb4ab"}.get(risk_level, "#64748b")
        risk_icon  = {"Normal": "✅", "Mild Risk": "⚠️", "High Risk": "🚨"}.get(risk_level, "")

        html += [
            "    <section id='interpretation'>",
            "      <h2>3. Cardiovascular Risk Assessment</h2>",
            f"      <div style='background:#f1f5f9; border-left: 8px solid {risk_color}; padding: 25px; border-radius: 8px; margin-bottom: 30px;'>",
            f"        <div style='font-size: 1.8rem; font-weight: 800; color: #0f172a;'>{risk_icon} {risk_level}</div>",
            f"        <div style='font-size: 1rem; color: #475569;'>Risk Score: <strong>{risk_res['score']:.0f}/100</strong> | Confidence: <strong>{risk_res['confidence']:.0f}%</strong></div>",
            "      </div>",
            "      <p><strong>Clinical Interpretation:</strong></p>",
            f"      <p style='background: #fff; border: 1px solid #e2e8f0; padding: 15px; border-radius: 6px;'>{interp.get('sdnn_class','—')}. {interp.get('autonomic','—')}. {interp.get('lf_hf_class','—')}.</p>",
            "      <h3>Risk Factor Breakdown</h3>"
        ]

        for metric, info in risk_res.get("flags", {}).items():
            status = info.get("status", "unavailable")
            m_color = {"normal": "#c3f400", "mild_risk": "#ffba38", "high_risk": "#f44336"}.get(status, "#64748b")
            m_icon  = {"normal": "✅", "mild_risk": "⚠️", "high_risk": "🚨"}.get(status, "—")
            html += [
                f"<div style='border-left: 4px solid {m_color}; padding: 10px 15px; background: #fff; border: 1px solid #f1f5f9; margin-bottom: 10px; border-radius: 4px;'>",
                f"  <div style='display:flex; justify-content:space-between; font-weight:600; font-size: 0.9rem;'><span>{m_icon} {metric}</span><span style='color:{m_color};'>{info.get('value','N/A')}</span></div>",
                f"  <div style='font-size: 0.8rem; color: #64748b;'>{info.get('clinical_note','')}</div>",
                "</div>"
            ]
        html.append("    </section>")

    html += [
        "    <footer>",
        "      <p>&copy; 2026 Clinical Sentinel Research Suite. All rights reserved.</p>",
        "      <p>This report is intended for research purposes only.</p>",
        "    </footer>",
        "  </div>",
        "</body>",
        "</html>"
    ]
    return "\n".join(html)

def build_docx_report(metrics_dict: dict, settings: dict, sqi_cache: dict) -> bytes:
    import docx
    from docx.shared import Inches, Pt
    import io
    doc = docx.Document()
    doc.add_heading("ECG & HRV Analysis Report", 0)
    
    for fname, m in metrics_dict.items():
        doc.add_heading(f"Analysis: {fname}", level=1)
        
        # Global Summary
        doc.add_heading("Global ECG Summary", level=2)
        doc.add_paragraph(f"Mean Heart Rate: {m.get('Mean HR (bpm)', 'N/A')} bpm")
        doc.add_paragraph(f"SDNN (Overall HRV): {_safe(m.get('SDNN (ms)'))} ms")
        doc.add_paragraph(f"RMSSD (Vagal Tone): {_safe(m.get('RMSSD (ms)'))} ms")
        
        charts = _generate_report_charts(fname)
        
        # 1. ECG Signal Section
        doc.add_heading("1. ECG Signal (Raw vs Filtered)", level=2)
        if "ecg_raw_filt" in charts:
            doc.add_picture(io.BytesIO(charts["ecg_raw_filt"]), width=Inches(6))

        # 2. R-Peak Detection
        doc.add_heading("2. R-Peak Detection Overlay", level=2)
        if "rpeaks" in charts:
            doc.add_picture(io.BytesIO(charts["rpeaks"]), width=Inches(6))

        # 3. RR Interval Analysis
        doc.add_heading("3. RR Interval Analysis & Tachogram", level=2)
        if "rr_tachogram" in charts:
            doc.add_picture(io.BytesIO(charts["rr_tachogram"]), width=Inches(6))

        # 4. Ectopic Correction
        doc.add_heading("4. Ectopic Beat Correction", level=2)
        if "ectopic_corr" in charts:
            doc.add_picture(io.BytesIO(charts["ectopic_corr"]), width=Inches(6))

        # 5. Frequency Domain
        doc.add_heading("5. Frequency Domain (PSD)", level=2)
        if "psd" in charts:
            doc.add_picture(io.BytesIO(charts["psd"]), width=Inches(6))

        # 6. Non-linear HRV
        doc.add_heading("6. Non-linear HRV (Poincaré & DFA)", level=2)
        if "poincare" in charts:
            doc.add_picture(io.BytesIO(charts["poincare"]), width=Inches(4))
        if "dfa" in charts:
            doc.add_picture(io.BytesIO(charts["dfa"]), width=Inches(6))

        # Tables
        doc.add_heading("Full Metrics Summary", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        metrics_to_list = [
            "Mean RR", "SDNN", "RMSSD", "NN50", "pNN50", 
            "LF Power", "HF Power", "LF/HF Ratio",
            "SD1", "SD2", "DFA α1", "Sample Entropy"
        ]
        for k in metrics_to_list:
            val = m.get(k) or m.get(k+" (ms)") or m.get(k+" (%)") or m.get(k+" (ms2)") or m.get(k+" (ms²)")
            row_cells = table.add_row().cells
            row_cells[0].text = k
            row_cells[1].text = _safe(val)

        # Interpretation
        doc.add_heading("Clinical Interpretation", level=2)
        interp = interpret_hrv(m, m)
        doc.add_paragraph(f"Analysis Status: {interp.get('sdnn_class','—')}")
        doc.add_paragraph(f"Vagal Activity: {interp.get('autonomic','—')}")
        doc.add_paragraph(f"Sympathovagal Balance: {interp.get('lf_hf_class','—')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    inject_stitch_theme()
    render_sidebar_settings()
    sentinel_header("Report Generation", badge="Export")
    pipeline_status_bar("HRV")

    metrics_dict = st.session_state.get("metrics",   {})
    sqi_cache    = st.session_state.get("sqi_cache", {})
    settings     = {k: st.session_state.get(k) for k in [
        "sfreq", "lowcut", "highcut", "filter_order", "rpeak_method", "remove_ectopic", "ectopic_method",
        "lf_min", "lf_max", "hf_min", "hf_max",
    ]}

    if not metrics_dict:
        st.warning("No analysis results found. Complete the pipeline first.")
        return

    section_header("Export Options")
    
    col_html = st.columns([1])[0]

    with col_html:
        if st.button("Publish (HTML)", use_container_width=True, help="Generate a self-contained interactive MATLAB-style report"):
            with st.spinner("Publishing to HTML..."):
                st.session_state["report_errors"] = []
                st.session_state["html_src"] = build_html_report(metrics_dict, settings, sqi_cache)
                if st.session_state.get("report_errors"):
                    for err in st.session_state["report_errors"]:
                        st.warning(f"⚠️ {err}")
                
        if "html_src" in st.session_state:
            st.download_button("Download Published HTML", data=st.session_state["html_src"], file_name="HRV_Published_Report.html", mime="text/html", use_container_width=True)

    with col_latex:
        if st.button("LaTeX Source", use_container_width=True):
            with st.spinner("Building LaTeX..."):
                st.session_state["latex_src"] = build_latex_report(metrics_dict, settings, sqi_cache)
                
        if "latex_src" in st.session_state:
            st.download_button("Download LaTeX", data=st.session_state["latex_src"], file_name="HRV_Report.tex", mime="text/plain", use_container_width=True)

    st.markdown("---")
    section_header("Clinical Interpretation & Risk Summary")

    RISK_COLORS = {"Normal": "#4caf7d", "Mild Risk": "#ffba38", "High Risk": "#f44336"}
    RISK_ICONS  = {"Normal": "✅", "Mild Risk": "⚠️", "High Risk": "🚨"}
    raw_rr_cache = st.session_state.get("raw_rr_intervals", {})

    for fname, m in metrics_dict.items():
        interp  = interpret_hrv(m, m)
        sqi     = sqi_cache.get(fname, {})
        lf_hf   = m.get("LF/HF Ratio", float('nan'))
        sqi_lbl = sqi.get("quality_label", "—") if sqi else "—"
        conf    = m.get("Confidence (%)", 100)

        # Compute risk
        raw_rr_f = raw_rr_cache.get(fname)
        pct_e = 0.0
        if raw_rr_f is not None and len(raw_rr_f) > 0:
            from utils.hrv_analysis import detect_ectopic_beats as _deb
            import numpy as _np
            _msk = _deb(raw_rr_f)
            pct_e = float(_np.sum(_msk)) / len(raw_rr_f) * 100
        
        # Pass SQI to risk classifier for adaptive confidence
        risk_res   = classify_cardiovascular_risk(m, pct_ectopic=pct_e, use_ml=True, sqi=sqi)
        risk_level = risk_res["risk_level"]
        
        # New Hero-style banner colors
        banner_styles = {
            "Normal":    {"color": "#c3f400", "bg": "#12140c", "border": "#c3f400", "icon": "✅"},
            "Mild Risk": {"color": "#ffba38", "bg": "#1a1608", "border": "#ffba38", "icon": "⚠️"},
            "High Risk": {"color": "#ffb4ab", "bg": "#1a0d0d", "border": "#ffb4ab", "icon": "🚨"},
        }
        bst = banner_styles.get(risk_level, banner_styles["Normal"])
        
        risk_score = risk_res["score"]
        risk_conf  = risk_res["confidence"]

        lf_hf_str = f"{lf_hf:.2f}" if isinstance(lf_hf, float) and lf_hf == lf_hf else "N/A"
        st.markdown(f"""
        <div style="background:{bst['bg']};border:1px solid {bst['border']};
                    border-left:4px solid {bst['border']};border-radius:0.4rem;
                    padding:1rem;margin-bottom:0.75rem;">
          <div style="display:flex;justify-content:space-between;margin-bottom:0.4rem;">
            <span style="font-weight:800;color:#849396;">{fname}</span>
            <div style="display:flex;gap:0.4rem;align-items:center;">
              <span style="background:{bst['color']};color:#000;font-size:0.6rem;
                           font-weight:800;padding:0.15rem 0.5rem;border-radius:0.2rem;">
                {bst['icon']} {risk_level} ({risk_score:.0f}/100)
              </span>
              <span style="background:#00daf3;color:#000;font-size:0.6rem;
                           padding:0.15rem 0.5rem;border-radius:0.2rem;font-weight:700;">SQI: {sqi_lbl}</span>
              <span style="background:#bac9cc;color:#000;font-size:0.6rem;
                           padding:0.15rem 0.5rem;border-radius:0.2rem;font-weight:700;">Confidence: {risk_conf:.0f}%</span>
            </div>
          </div>
          <div style="font-size:0.8rem;color:#bac9cc;">
            <strong>HRV Status:</strong> {interp.get('sdnn_class','—')}<br>
            <strong>Vagal Tone:</strong> {interp.get('autonomic','—')}<br>
            <strong>Sympathovagal (LF/HF={lf_hf_str}):</strong> {interp.get('lf_hf_class','—')}
          </div>
        </div>""", unsafe_allow_html=True)

    # ── Heart disease section header ───────────────────────────────────────────
    st.markdown("---")
    section_header("Heart Disease Risk Assessment (All Files)")
    st.info("For detailed per-metric breakdown and ML confidence scores, visit **Dashboard 10 · Heart Disease Detection**.")


if __name__ == "__main__":
    main()
